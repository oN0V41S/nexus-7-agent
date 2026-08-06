"""
Módulo de Geração Contextualizada (REQ-005).

Gera currículo adaptado à vaga (DOCX + PDF + MD) e carta de apresentação (TXT)
usando Ollama para conteúdo contextualizado com fallback para templates.

v4.0.0: Markdown-first resume pipeline — gera .md e converte para .docx e .pdf
         com formatação rica (negrito, itálico, hyperlinks, bullet points).
         PDF com limite de 2 páginas via fpdf2.
"""
import datetime
import json
import logging
import re
import warnings
from pathlib import Path
from typing import Optional

from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement

from fpdf import FPDF

from src.job_apply_agent.config import OLLAMA_URL, OLLAMA_MODEL, PROFILE_DIR
from src.job_apply_agent.analyzer import _call_ollama

logger = logging.getLogger(__name__)


# ─── Orçamento de tamanho ATS (critérios de 1-2 páginas) ──────────────────────
# Currículos devem ser condensados e direcionados à vaga. Estes limites evitam
# o estouro para 3+ páginas incluindo todo o histórico do candidato.
MAX_SKILLS = 8           # máx. de skills relevantes exibidas
MAX_CERTS = 2            # máx. de certificações relevantes
MAX_PROJECTS = 1         # máx. de projetos exibidos
MAX_BULLETS_PER_ROLE = 2 # máx. de bullets por experiência
MAX_ROLES = 2            # máx. de experiências (mantém as mais relevantes)

_STOPWORDS = {
    "de", "da", "do", "das", "dos", "e", "ou", "com", "para", "em", "na", "no",
    "nas", "nos", "por", "que", "a", "o", "as", "os", "um", "uma", "você", "se",
    "the", "and", "with", "for", "to", "of", "in", "on", "at", "is", "are", "via",
    "using", "use", "como", "ser", "são", "foi", "tem", "suas", "seus", "candidato",
}


def _extract_job_keywords(job: dict) -> set[str]:
    """
    Extrai conjunto de tokens relevantes da vaga para pontuar relevância.

    Combina título, descrição, requisitos obrigatórios, stack e strengths, ignorando
    stopwords e tokens curtos.
    """
    texts: list[str] = []
    texts.append(str(job.get("title", "")))
    texts.append(str(job.get("description", "")))
    reqs = job.get("requirements")
    if isinstance(reqs, dict):
        texts.extend(str(r) for r in reqs.get("mandatory", []))
        texts.extend(str(r) for r in reqs.get("nice_to_have", []))
    elif isinstance(reqs, (list, tuple)):
        texts.extend(str(r) for r in reqs)
    # Adiciona stack da vaga (importante para filtering)
    stack = job.get("stack", [])
    if isinstance(stack, (list, tuple)):
        texts.extend(str(s) for s in stack)
    texts.extend(str(s) for s in job.get("strengths", []))
    texts.extend(str(g) for g in job.get("gaps", []))

    tokens: set[str] = set()
    for text in texts:
        for tok in re.split(r"[^a-zA-ZÀ-ÿ0-9+#.]+", text.lower()):
            tok = tok.strip()
            if len(tok) >= 3 and tok not in _STOPWORDS:
                tokens.add(tok)
    return tokens


def _relevance_score(text: str, keywords: set[str]) -> int:
    """Conta quantos keywords da vaga aparecem no texto informado."""
    if not keywords:
        return 0
    lowered = text.lower()
    return sum(1 for kw in keywords if kw in lowered)


def _filter_by_relevance(
    items: list, keywords: set[str], max_n: int, text_fn=lambda x: str(x)
) -> list:
    """
    Retorna os ``max_n`` itens mais relevantes para a vaga.

    Itens com mesma pontuação preservam a ordem original (estável). Se não
    houver keywords, retorna os ``max_n`` primeiros.
    """
    if not items:
        return []
    if not keywords:
        return items[:max_n]
    scored = sorted(
        items, key=lambda it: _relevance_score(text_fn(it), keywords), reverse=True
    )
    return scored[:max_n]


# ─── Templates de carta ──────────────────────────────────────────────────────

COVER_LETTER_TEMPLATE = """{candidate_name}
{candidate_email} | {candidate_phone} | {candidate_location}

{date}

{company_name}
{company_location}

Ref.: Candidatura à vaga de {job_title}

Prezado(a) Recrutador(a),

{body}

Atenciosamente,
{candidate_name}
"""


COVER_LETTER_BODY_TEMPLATE = (
    "Gostaria de expressar meu interesse pela vaga de {job_title} na {company_name}. "
    "Com base na descrição da vaga, identifico forte alinhamento entre minha trajetória "
    "profissional e os requisitos apresentados.\n\n"
    "{strengths_paragraph}\n\n"
    "Estou entusiasmado(a) com a possibilidade de contribuir para os desafios da "
    "{company_name} e acredito que minha experiência pode agregar valor ao time.\n\n"
    "Coloco-me à disposição para uma conversa aprofundada sobre minha candidatura."
)


# ─── Geração via LLM ─────────────────────────────────────────────────────────


def _build_ollama_prompt(section: str, profile: dict, job: dict) -> str:
    """Constrói prompt para o Ollama gerar conteúdo contextualizado."""
    prompts = {
        "resume_summary": (
            "Gere um resumo profissional de 3-4 linhas para um currículo adaptado "
            "à seguinte vaga. Destaque as experiências e habilidades mais relevantes "
            "para a posição. Seja conciso e impactante.\n\n"
            f"Perfil do candidato:\n{json.dumps(profile.get('skills', []), ensure_ascii=False)}\n"
            f"Experiência: {profile.get('experience', '')[:500]}\n\n"
            f"Vaga: {job.get('title', '')} em {job.get('company', '')}\n"
            f"Descrição: {job.get('description', '')[:1000]}\n\n"
            "Resumo profissional (texto puro, sem markdown):"
        ),
        "cover_letter": (
            "Gere o corpo de uma carta de apresentação profissional (2-3 parágrafos) "
            "para a candidatura à vaga abaixo. Use tom profissional e entusiasmado. "
            "Destaque skills relevantes. Mencione contribuições específicas.\n\n"
            f"Candidato: {json.dumps(profile.get('skills', []), ensure_ascii=False)}\n"
            f"Experiência: {profile.get('experience', '')[:800]}\n\n"
            f"Vaga: {job.get('title', '')} em {job.get('company', '')}\n"
            f"Descrição: {job.get('description', '')[:1000]}\n\n"
            "Texto da carta (apenas o corpo, sem saudações, sem markdown):"
        ),
    }
    return prompts.get(section, "")


def _generate_text_ollama(prompt: str) -> Optional[str]:
    """Gera texto via Ollama com tratamento de erro."""
    result = _call_ollama(prompt)
    if result:
        result = re.sub(r"```[\w]*\n?", "", result).strip()
        return result
    return None


# ─── Helpers de formatação DOCX ──────────────────────────────────────────────


def _add_hyperlink(paragraph: Paragraph, text: str, url: str):
    """Add a hyperlink to a paragraph (python-docx).

    Falls back to plain text if URL is empty, not a string, or does not
    start with http://, https://, or mailto:.
    """
    if not url or not isinstance(url, str) or not url.startswith(('http://', 'https://', 'mailto:')):
        run = paragraph.add_run(text)
        run.font.size = Pt(10.5)
        run.font.name = "Calibri"
        return paragraph

    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    # Cor azul + sublinhado (Hyperlink style pode nao existir no template)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return paragraph


def _parse_inline_md(paragraph: Paragraph, text: str, base_size=Pt(10.5)):
    """Parse inline Markdown (``**bold**``, ``[text](url)``) into formatted runs."""
    pattern = r"\*\*(.+?)\*\*|\[(.+?)\]\((.+?)\)"
    last_end = 0
    for match in re.finditer(pattern, text):
        # Plain text before this match
        if match.start() > last_end:
            run = paragraph.add_run(text[last_end : match.start()])
            run.font.size = base_size
            run.font.name = "Calibri"
        # Matched token
        if match.group(1):  # **bold**
            run = paragraph.add_run(match.group(1))
            run.bold = True
            run.font.size = base_size
            run.font.name = "Calibri"
        elif match.group(2):  # [text](url)
            _add_hyperlink(paragraph, match.group(2), match.group(3))
        last_end = match.end()
    # Remaining plain text
    if last_end < len(text):
        run = paragraph.add_run(text[last_end:])
        run.font.size = base_size
        run.font.name = "Calibri"


def _setup_doc(doc: Document) -> Document:
    """Apply default formatting to a Document."""
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.0
    return doc


# ─── Parsing de seções do perfil ─────────────────────────────────────────────


def _build_section_list(profile: dict, job: dict) -> list[dict]:
    """
    Analisa quais seções do perfil devem aparecer com base nos requisitos da vaga.

    Returns:
        Lista de dicts::
            [{"name": "Resumo Profissional", "key": "summary",
              "required": True, "order": 1}, ...]
    """
    sections = []
    order = 0

    # ── Sempre presentes ──
    order += 1
    sections.append({
        "name": "Resumo Profissional",
        "key": "summary",
        "required": True,
        "order": order,
    })
    order += 1
    sections.append({
        "name": "Habilidades Técnicas",
        "key": "skills",
        "required": True,
        "order": order,
    })
    order += 1
    sections.append({
        "name": "Experiência Profissional",
        "key": "experience",
        "required": True,
        "order": order,
    })
    order += 1
    sections.append({
        "name": "Formação Acadêmica",
        "key": "education",
        "required": True,
        "order": order,
    })

    # ── Condicionais ──
    if profile.get("languages"):
        order += 1
        sections.append({
            "name": "Idiomas",
            "key": "languages",
            "required": False,
            "order": order,
        })
    if profile.get("certifications"):
        order += 1
        sections.append({
            "name": "Certificações",
            "key": "certifications",
            "required": False,
            "order": order,
        })
    if profile.get("projects"):
        order += 1
        sections.append({
            "name": "Projetos",
            "key": "projects",
            "required": False,
            "order": order,
        })

    # ── Links condicionais ──
    profile_has_links = any(
        profile.get(k) for k in ("github", "linkedin", "portfolio")
    )
    if profile_has_links:
        order += 1
        sections.append({
            "name": "Links",
            "key": "links",
            "required": False,
            "order": order,
        })

    return sections


# ─── Parsing de experiência / educação ───────────────────────────────────────


def _parse_experience_entries(exp_text: str) -> list[dict]:
    """
    Parseia texto de experiência em entries estruturadas.

    Formato esperado (separador ``. `` entre entries)::

        Cargo na Empresa (Período): Responsabilidades.
        Cargo na Empresa (Período): Responsabilidades.

    Returns:
        Lista de dicts com as chaves ``role``, ``company``, ``period``, ``details``.
    """
    entries = []
    if not exp_text:
        return entries

    parts = re.split(r"(?<=\.)\s+(?=[A-ZÀ-Ú])", exp_text)
    for part in parts:
        part = part.strip().rstrip(".")
        if not part:
            continue
        m = re.match(
            r"^(.+)\s+(?:na|no|em|–|-)\s+(.+?)\s*\(([^)]+)\)(?::\s*(.*))?$",
            part,
            re.UNICODE,
        )
        if m:
            entries.append({
                "role": m.group(1).strip(),
                "company": m.group(2).strip(),
                "period": m.group(3).strip(),
                "details": (m.group(4) or "").strip(),
            })
        else:
            # Fallback — entry não-parseada vira detalhe sem role
            logger.debug("Experiência não-parseada no formato esperado: %s", part[:80])
            entries.append({
                "role": "",
                "company": "",
                "period": "",
                "details": part,
            })
    return entries


def _parse_education_entries(edu_text: str) -> list[dict]:
    """
    Parseia texto de formação em entries estruturadas.

    Formato::

        Curso na Instituição (Período). Curso na Instituição (Período).

    Returns:
        Lista de dicts com ``course``, ``institution``, ``period``.
    """
    entries = []
    if not edu_text:
        return entries

    parts = re.split(r"(?<=\.)\s+(?=[A-ZÀ-Ú])", edu_text)
    for part in parts:
        part = part.strip().rstrip(".")
        if not part:
            continue
        m = re.match(
            r"^(.+)\s+(?:na|no|em|–|-)\s+(.+?)\s*\(([^)]+)\)\s*$",
            part,
            re.UNICODE,
        )
        if m:
            entries.append({
                "course": m.group(1).strip(),
                "institution": m.group(2).strip(),
                "period": m.group(3).strip(),
            })
        else:
            logger.debug("Formação não-parseada no formato esperado: %s", part[:80])
            entries.append({
                "course": part,
                "institution": "",
                "period": "",
            })
    return entries




# ─── Fallback inteligente para resumo adaptado ─────────────────────────────────


def _build_smart_summary(profile: dict, job: dict) -> str:
    """
    Gera resumo profissional adaptado à vaga sem Ollama.

    Usa gaps/strengths da análise heurística para enfatizar skills relevantes.
    Retorna string com resumo contextualizado.
    """
    base_summary = profile.get("summary", "")
    strengths = job.get("strengths", [])
    gaps = job.get("gaps", [])
    job_title = job.get("title", "")

    # Se não há strengths específicos, retorna o summary original
    if not strengths:
        return base_summary

    # Constrói adapted_summary: adiciona ênfase nas strengths.
    # Usa o texto original das strengths (já são frases legíveis) em vez de
    # .title(), que quebra a capitalização de termos técnicos (ex: "Node.js").
    strengths_text = "; ".join(s.strip() for s in strengths[:4])

    # Identifica o "foco" da vaga a partir do título
    focus_areas = {
        "infraestrutura": "infraestrutura de TI",
        "suporte": "suporte e infraestrutura",
        "dados": "análise de dados",
        "data": "análise de dados",
        "engenheiro de dados": "engenharia de dados",
        "bi": "business intelligence",
        "analytics": "business intelligence e analytics",
        "devops": "infraestrutura e automação DevOps",
        "devsecops": "segurança e automação DevOps",
        "java": "desenvolvimento backend Java",
        "frontend": "desenvolvimento frontend",
        "full stack": "desenvolvimento full stack",
        "ia": "inteligência artificial aplicada",
        "machine learning": "machine learning",
        "segurança": "segurança da informação",
        "cyber": "cibersegurança",
    }

    focus = ""
    title_lower = job_title.lower()
    for keyword, area in focus_areas.items():
        if keyword in title_lower:
            focus = area
            break

    if focus:
        adapted = (
            f"{profile.get('name', 'Profissional')} com experiência em {focus}. "
            f"Principais competências: {strengths_text}. "
        )
        # Adiciona a formação se disponível
        edu = profile.get("education", "")
        if edu:
            edu_short = edu.split(".")[0].strip() if "." in edu else edu[:60]
            adapted += f"Formação em {edu_short}. "
        # Adiciona o summary original como complemento se não for redundante
        if base_summary and not any(s.lower() in base_summary.lower() for s in strengths):
            adapted += base_summary
        return adapted.strip()

    # Fallback: se não achou foco, usa base_summary + strengths de forma fluida
    if base_summary:
        return f"{base_summary.rstrip('.')}. Principais competências alinhadas à vaga: {strengths_text}."
    return f"Profissional de TI com experiência em {strengths_text}."


# ─── Geração Markdown do currículo ───────────────────────────────────────────


def _build_resume_markdown(profile: dict, job: dict) -> str:
    """
    Gera uma string Markdown completa para o currículo adaptado.

    Seções são definidas por ``_build_section_list``.  O summary é gerado via
    Ollama quando disponível (fallback para ``profile["summary"]``).
    """
    lines = []
    name = profile.get("name", "Candidato(a)")
    email = profile.get("email", "")
    phone = profile.get("phone", "")
    location = profile.get("location", "")

    # ── Cabeçalho ────────────────────────────────────────────────────────
    lines.append(f"# {name}")
    parts_filtered = [p for p in (email, phone, location) if p]
    lines.append(" | ".join(parts_filtered))
    lines.append("")

    sections = _build_section_list(profile, job)

    for sec in sections:
        key = sec["key"]
        # ── Resumo Profissional ──────────────────────────────────────────
        if key == "summary":
            lines.append(f"## {sec['name']}")
            summary = _generate_text_ollama(
                _build_ollama_prompt("resume_summary", profile, job)
            )
            if not summary:
                # Fallback inteligente: usa _build_smart_summary
                summary = _build_smart_summary(profile, job)
            if summary:
                lines.append(summary)
            lines.append("")

        # ── Habilidades Técnicas ─────────────────────────────────────────
        elif key == "skills":
            skills = profile.get("skills", [])
            if skills:
                # Mantém apenas as skills mais relevantes para a vaga (ATS: 1-2 páginas)
                job_kw = _extract_job_keywords(job)
                relevant = _filter_by_relevance(skills, job_kw, MAX_SKILLS)
                # Preserva a ordem original do perfil entre as selecionadas
                seen = set(id(s) for s in relevant)
                ordered = [s for s in skills if id(s) in seen][:MAX_SKILLS]
                if not ordered:
                    ordered = skills[:MAX_SKILLS]
                lines.append(f"## {sec['name']}")
                grouped = [ordered[i : i + 5] for i in range(0, len(ordered), 5)]
                for group in grouped:
                    lines.append(f"- {', '.join(group)}")
                lines.append("")

        # ── Experiência Profissional ──────────────────────────────────────
        elif key == "experience":
            # PRIORIDADE: usar experience_raw (com bullets preservados)
            exp_raw = profile.get("experience_raw", "")
            exp_text = profile.get("experience", "")
            job_kw = _extract_job_keywords(job)

            def _render_role(header: str, bullets: list[str]) -> None:
                """Renderiza um cargo com cabeçalho e no máx. MAX_BULLETS_PER_ROLE bullets."""
                if header:
                    m = re.match(
                        r"^(.+?)\s+(?:na|no|em|–|-)\s+(.+?)\s*\(([^)]+)\)\s*$",
                        header, re.UNICODE,
                    )
                    if m:
                        lines.append(
                            f"### **{m.group(1).strip()}** | {m.group(2).strip()} | {m.group(3).strip()}"
                        )
                    else:
                        lines.append(f"### {header}")
                for detail_line in bullets[:MAX_BULLETS_PER_ROLE]:
                    detail_line = detail_line.strip().lstrip("•- \t")
                    if detail_line:
                        lines.append(f"- {detail_line}")

            if exp_raw and "•" in exp_raw:
                # Parse do formato raw com bullets
                blocks = exp_raw.strip().split("\n\n")
                parsed = []
                for block in blocks:
                    block = block.strip()
                    if not block:
                        continue
                    lines_in_block = block.split("\n")
                    header = lines_in_block[0].strip()
                    bullets = [
                        b.strip().lstrip("•- \t")
                        for b in lines_in_block[1:]
                        if b.strip().lstrip("•- \t")
                    ]
                    parsed.append((header, bullets))
                # Mantém apenas os MAX_ROLES cargos mais relevantes para a vaga
                if len(parsed) > MAX_ROLES:
                    parsed = _filter_by_relevance(
                        parsed, job_kw, MAX_ROLES,
                        text_fn=lambda pb: f"{pb[0]} {' '.join(pb[1])}",
                    )
                if parsed:
                    lines.append(f"## {sec['name']}")
                    for header, bullets in parsed:
                        _render_role(header, bullets)
                    lines.append("")
            elif exp_text:
                # Fallback: usa o texto normalizado
                entries = _parse_experience_entries(exp_text)
                if len(entries) > MAX_ROLES:
                    entries = _filter_by_relevance(
                        entries, job_kw, MAX_ROLES,
                        text_fn=lambda e: f"{e.get('role','')} {e.get('details','')}",
                    )
                if entries:
                    lines.append(f"## {sec['name']}")
                    for entry in entries:
                        role = entry.get("role", "")
                        company = entry.get("company", "")
                        period = entry.get("period", "")
                        details = entry.get("details", "")
                        header = ""
                        if role and company:
                            header = f"{role} | {company} | {period}"
                        elif role:
                            header = f"{role} | {period}"
                        bullets = []
                        if details:
                            items = re.split(r"\.\s+(?=[A-ZÀ-Ú])", details)
                            bullets = [i.strip().rstrip(".") for i in items if i.strip()]
                        _render_role(header, bullets)
                    lines.append("")

        # ── Formação Acadêmica ────────────────────────────────────────────
        elif key == "education":
            edu_text = profile.get("education", "")
            if edu_text:
                lines.append(f"## {sec['name']}")
                entries = _parse_education_entries(edu_text)
                for entry in entries:
                    course = entry.get("course", "")
                    institution = entry.get("institution", "")
                    period = entry.get("period", "")
                    parts_line = [f"**{course}**"]
                    if institution:
                        parts_line.append(institution)
                    if period:
                        parts_line.append(period)
                    lines.append(f"### {' | '.join(parts_line)}")
                lines.append("")

        # ── Idiomas ──────────────────────────────────────────────────────
        elif key == "languages":
            langs = profile.get("languages", "")
            lines.append(f"## {sec['name']}")
            if isinstance(langs, str):
                for lang in langs.split(","):
                    lang = lang.strip()
                    if lang:
                        lines.append(f"- {lang}")
            elif isinstance(langs, list):
                for lang in langs:
                    lines.append(f"- {lang}")
            lines.append("")

        # ── Certificações ────────────────────────────────────────────────
        elif key == "certifications":
            certs = profile.get("certifications", [])
            if isinstance(certs, list):
                # Mantém apenas as MAX_CERTS mais relevantes para a vaga
                job_kw = _extract_job_keywords(job)
                certs = _filter_by_relevance(certs, job_kw, MAX_CERTS, text_fn=str)
            else:
                certs = [certs]
            if certs:
                lines.append(f"## {sec['name']}")
                for c in certs:
                    lines.append(f"- {c}")
                lines.append("")

        # ── Projetos ─────────────────────────────────────────────────────
        elif key == "projects":
            projs = profile.get("projects", [])
            if isinstance(projs, list):
                # Mantém apenas os MAX_PROJECTS mais relevantes para a vaga
                job_kw = _extract_job_keywords(job)
                projs = _filter_by_relevance(
                    projs, job_kw, MAX_PROJECTS,
                    text_fn=lambda p: (
                        f"{p.get('name','')} {p.get('description','')}"
                        if isinstance(p, dict) else str(p)
                    ),
                )
            else:
                projs = [projs]
            if projs:
                lines.append(f"## {sec['name']}")
                for p in projs:
                    if isinstance(p, dict):
                        name_p = p.get("name", "")
                        desc_p = p.get("description", "")
                        link_p = p.get("link", "")
                        # Nome do projeto como H3
                        lines.append(f"### **{name_p}**")
                        # Descrição completa como bullet
                        if desc_p:
                            lines.append(f"- {desc_p}")
                        # Link se existir
                        if link_p:
                            lines.append(f"- Link: {link_p}")
                    else:
                        lines.append(f"- {p}")
                lines.append("")

        # ── Links ────────────────────────────────────────────────────────
        elif key == "links":
            lines.append(f"## {sec['name']}")
            github = profile.get("github", "")
            linkedin = profile.get("linkedin", "")
            portfolio = profile.get("portfolio", "")
            if github:
                lines.append(f"- GitHub: {github}")
            if linkedin:
                lines.append(f"- LinkedIn: {linkedin}")
            if portfolio:
                lines.append(f"- Portfolio: {portfolio}")
            lines.append("")

    return "\n".join(lines)


# ─── Conversão Markdown → DOCX ──────────────────────────────────────────────


def _md_to_docx(md_text: str) -> Document:
    """
    Converte string Markdown em um Document python-docx formatado.

    Parseia cabeçalhos (``#`` / ``##`` / ``###``), **negrito**,
    ``[links](url)``, bullet points (``- ``), e linhas de contato.
    """
    doc = Document()
    _setup_doc(doc)

    lines = md_text.split("\n")
    i = 0
    after_h1 = False  # next non-empty line after # heading = contact info

    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        # ── Linha vazia ──────────────────────────────────────────────────
        if not stripped:
            p = doc.add_paragraph("")
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # ── H1 — Nome ────────────────────────────────────────────────────
        if stripped.startswith("# ") and not stripped.startswith("## "):
            text = stripped[2:]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = "Calibri"
            p.paragraph_format.space_after = Pt(4)
            after_h1 = True
            i += 1
            continue

        # ── H2 — Seção ───────────────────────────────────────────────────
        if stripped.startswith("## ") and not stripped.startswith("### "):
            text = stripped[3:]
            p = doc.add_paragraph()
            run = p.add_run(text.upper())
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = "Calibri"
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            after_h1 = False
            i += 1
            continue

        # ── H3 — Sub-seção (experiência / formação) ──────────────────────
        if stripped.startswith("### "):
            text = stripped[4:]
            p = doc.add_paragraph()
            _parse_inline_md(p, text, base_size=Pt(10.5))
            # Toda a H3 é itálica (além do bold das partes **)
            for run in p.runs:
                run.italic = True
                run.font.name = "Calibri"
                if not run.bold:
                    run.font.size = Pt(10.5)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(0)
            after_h1 = False
            i += 1
            continue

        # ── Linha de contato (após H1 ou contém @) ───────────────────────
        if after_h1 or ("@" in stripped and "|" in stripped):
            p = doc.add_paragraph()
            _parse_inline_md(p, stripped, base_size=Pt(10))
            p.paragraph_format.space_after = Pt(4)
            after_h1 = False
            i += 1
            continue

        # ── Bullet point ─────────────────────────────────────────────────
        if stripped.startswith("- "):
            text = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            bullet = p.add_run("• ")
            bullet.font.size = Pt(10.5)
            bullet.font.name = "Calibri"
            _parse_inline_md(p, text, base_size=Pt(10.5))
            i += 1
            continue

        # ── Parágrafo normal ─────────────────────────────────────────────
        p = doc.add_paragraph()
        _parse_inline_md(p, stripped, base_size=Pt(10.5))
        p.paragraph_format.space_after = Pt(2)
        after_h1 = False
        i += 1

    return doc


# ─── Conversão Markdown → PDF ───────────────────────────────────────────────


class _ResumePDF(FPDF):
    """Classe PDF personalizada para currículos com formatação ATS-friendly."""

    def __init__(self):
        super().__init__()
        self.set_auto_page_break(auto=True, margin=15)
        # Adiciona fonte Unicode para caracteres especiais
        self.add_font("DejaVu", "", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", uni=True)
        self.add_font("DejaVu", "B", "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", uni=True)
        # Usa fonte regular para itálico (fallback seguro)
        self.add_font("DejaVu", "I", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", uni=True)
        self.add_font("DejaVu", "BI", "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", uni=True)

    def header(self):
        pass

    def footer(self):
        self.set_y(-15)
        self.set_font("DejaVu", "I", 8)
        self.set_text_color(128, 128, 128)
        self.cell(0, 10, f"Página {self.page_no()}/{{nb}}", 0, 0, "C")


def _md_to_pdf(md_text: str, output_path: Path, max_pages: int = 2) -> Path:
    """
    Converte string Markdown em PDF formatado para ATS.

    Args:
        md_text: Conteúdo Markdown do currículo
        output_path: Caminho do arquivo PDF de saída
        max_pages: Número máximo de páginas (padrão: 2)

    Returns:
        Path do arquivo PDF gerado
    """
    pdf = _ResumePDF()
    pdf.alias_nb_pages()
    pdf.add_page()
    pdf.set_margins(15, 15, 15)

    lines = md_text.split("\n")
    i = 0
    after_h1 = False

    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        # Verifica se excedeu o limite de páginas
        if pdf.page_no() > max_pages:
            break

        # ── Linha vazia ──────────────────────────────────────────────────
        if not stripped:
            pdf.ln(2)
            i += 1
            continue

        # ── H1 — Nome ────────────────────────────────────────────────────
        if stripped.startswith("# ") and not stripped.startswith("## "):
            text = stripped[2:]
            pdf.set_font("DejaVu", "B", 14)
            pdf.set_text_color(0, 0, 0)
            pdf.multi_cell(0, 8, text, 0, "C")
            pdf.ln(2)
            after_h1 = True
            i += 1
            continue

        # ── H2 — Seção ───────────────────────────────────────────────────
        if stripped.startswith("## ") and not stripped.startswith("### "):
            text = stripped[3:].upper()
            pdf.set_font("DejaVu", "B", 11)
            pdf.set_text_color(0, 0, 0)
            pdf.ln(4)
            pdf.multi_cell(0, 6, text, 0, "L")
            pdf.line(15, pdf.get_y(), 195, pdf.get_y())
            pdf.ln(2)
            after_h1 = False
            i += 1
            continue

        # ── H3 — Sub-seção ───────────────────────────────────────────────
        if stripped.startswith("### "):
            text = stripped[4:]
            pdf.set_font("DejaVu", "BI", 10)
            pdf.set_text_color(0, 0, 0)
            # Remove formatação Markdown para PDF
            clean_text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            pdf.multi_cell(0, 5, clean_text, 0, "L")
            pdf.ln(1)
            after_h1 = False
            i += 1
            continue

        # ── Linha de contato (após H1) ───────────────────────────────────
        if after_h1 or ("@" in stripped and "|" in stripped):
            pdf.set_font("DejaVu", "", 9)
            pdf.set_text_color(80, 80, 80)
            # Remove formatação Markdown para PDF
            clean_text = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
            clean_text = re.sub(r"\[(.+?)\]\((.+?)\)", r"\1", clean_text)
            pdf.multi_cell(0, 4, clean_text, 0, "C")
            pdf.ln(2)
            after_h1 = False
            i += 1
            continue

        # ── Bullet point ─────────────────────────────────────────────────
        if stripped.startswith("- "):
            text = stripped[2:]
            pdf.set_font("DejaVu", "", 9.5)
            pdf.set_text_color(0, 0, 0)
            # Remove formatação Markdown para PDF
            clean_text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            clean_text = re.sub(r"\[(.+?)\]\((.+?)\)", r"\1", clean_text)
            # Usa multi_cell para bullets com wrap adequado
            bullet_text = f"• {clean_text}"
            pdf.set_x(20)  # Indentação para bullets
            pdf.multi_cell(175, 4, bullet_text, 0, "L")
            pdf.ln(1)
            i += 1
            continue

        # ── Parágrafo normal ─────────────────────────────────────────────
        pdf.set_font("DejaVu", "", 9.5)
        pdf.set_text_color(0, 0, 0)
        # Remove formatação Markdown e renderiza com wrap
        clean_text = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
        clean_text = re.sub(r"\[(.+?)\]\((.+?)\)", r"\1", clean_text)
        pdf.multi_cell(0, 4, clean_text, 0, "L")
        pdf.ln(2)
        after_h1 = False
        i += 1

    pdf.output(str(output_path))
    return output_path


def _render_pdf_inline(pdf: FPDF, text: str, base_size: float):
    """
    Renderiza texto com formatação inline (**negrito**) no PDF.

    Remove links [text](url) e mantém apenas o texto.
    """
    # Remove links e mantém apenas texto
    text = re.sub(r"\[(.+?)\]\((.+?)\)", r"\1", text)
    # Remove formatação de negrito para PDF (fpdf2 não suporta inline bold facilmente)
    clean_text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    pdf.multi_cell(0, 4, clean_text, 0, "L")


# ─── Geração de currículo adaptado (DOCX + PDF + MD) ────────────────────────


def _content_to_docx(text: str, title: str = "Resume") -> Document:
    """
    Renderiza texto em DOCX formatado (ATS-friendly) — **deprecated**.

    Mantida apenas para compatibilidade retroativa.  Use ``_md_to_docx`` para
    currículos novos com formatação Markdown.
    """
    warnings.warn(
        "_content_to_docx is deprecated, use _md_to_docx instead",
        DeprecationWarning,
        stacklevel=2,
    )
    doc = Document()
    _setup_doc(doc)

    if not text or not text.strip():
        text = "Documento gerado pelo Job Application Agent"

    for line in text.split("\n"):
        if not line.strip():
            p = doc.add_paragraph("")
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0

        run = p.add_run(line)

        if line.isupper() and len(line) > 3:
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = "Calibri"
        else:
            run.font.size = Pt(10.5)
            run.font.name = "Calibri"

    return doc


def generate_adapted_resume(profile: dict, job: dict, output_dir: Path, max_pages: int = 2) -> dict:
    """
    Gera currículo adaptado à vaga nos formatos MD, DOCX e PDF.

    Args:
        profile: Perfil do candidato (profile.json)
        job: Vaga alvo (com score, gaps, strengths)
        output_dir: Diretório de saída
        max_pages: Número máximo de páginas para o PDF (padrão: 2)

    Returns:
        dict com ``docx_path``, ``pdf_path`` e ``md_path``.
    """
    md_text = _build_resume_markdown(profile, job)

    # Valida completude
    warnings = validate_resume_completeness(md_text, profile)
    if warnings:
        logger.warning("⚠️ Currículo pode estar incompleto:")
        for w in warnings:
            logger.warning(f"   - {w}")

    # Salva Markdown
    md_path = output_dir / "resume_adapted.md"
    md_path.write_text(md_text, encoding="utf-8")
    logger.info(f"Currículo adaptado (MD) salvo: {md_path}")

    # Converte para DOCX
    doc = _md_to_docx(md_text)
    docx_path = output_dir / "resume_adapted.docx"
    doc.save(str(docx_path))
    logger.info(f"Currículo adaptado (DOCX) salvo: {docx_path}")

    # Converte para PDF (com limite de páginas)
    pdf_path = output_dir / "resume_adapted.pdf"
    _md_to_pdf(md_text, pdf_path, max_pages=max_pages)
    logger.info(f"Currículo adaptado (PDF) salvo: {pdf_path}")

    return {"docx_path": docx_path, "pdf_path": pdf_path, "md_path": md_path}


# ─── Geração de carta de apresentação (TXT) ──────────────────────────────────


def generate_cover_letter(profile: dict, job: dict, output_dir: Path) -> Path:
    """
    Gera TXT da carta de apresentação adaptada à vaga.

    Usa Ollama para gerar conteúdo; fallback para template.
    """
    company = job.get("company", "Empresa")
    title = job.get("title", "Vaga")
    strengths = job.get("strengths", [])

    # Tenta gerar corpo via Ollama
    body = _generate_text_ollama(_build_ollama_prompt("cover_letter", profile, job))

    # Fallback para template
    if not body:
        strengths_text = (
            f"Entre minhas principais qualificações, destaco: "
            f"{', '.join(strengths)}."
        ) if strengths else (
            "Minha trajetória profissional e habilidades técnicas estão alinhadas "
            "com os requisitos da posição."
        )
        body = COVER_LETTER_BODY_TEMPLATE.format(
            job_title=title,
            company_name=company,
            strengths_paragraph=strengths_text,
        )

    letter = COVER_LETTER_TEMPLATE.format(
        candidate_name=profile.get("name", "Candidato(a)"),
        candidate_email=profile.get("email", "email@exemplo.com"),
        candidate_phone=profile.get("phone", "(11) 99999-9999"),
        candidate_location=profile.get("location", "São Paulo, SP"),
        date=datetime.date.today().strftime("%d/%m/%Y"),
        company_name=company,
        company_location=job.get("location", ""),
        job_title=title,
        body=body,
    )

    output_path = output_dir / "cover_letter.txt"
    output_path.write_text(letter, encoding="utf-8")
    logger.info(f"Carta de apresentação salva: {output_path}")
    return output_path


# ─── Orquestração ────────────────────────────────────────────────────────────


def generate_application(profile: dict, job: dict, output_dir: Path, max_pages: int = 2) -> dict:
    """
    Pipeline completo de geração: currículo adaptado (MD + DOCX + PDF) + carta (TXT).

    Args:
        profile: Perfil do candidato
        job: Vaga alvo (enriquecida com análise)
        output_dir: Diretório de saída
        max_pages: Número máximo de páginas para o PDF (padrão: 2)

    Returns:
        dict com ``resume_path`` (DOCX, compatibilidade), ``resume_docx``,
        ``resume_pdf``, ``resume_md`` e ``cover_letter_path``.
    """
    resume_result = generate_adapted_resume(profile, job, output_dir, max_pages=max_pages)
    letter_path = generate_cover_letter(profile, job, output_dir)

    return {
        "resume_path": str(resume_result["docx_path"]),
        "resume_docx": str(resume_result["docx_path"]),
        "resume_pdf": str(resume_result["pdf_path"]),
        "resume_md": str(resume_result["md_path"]),
        "cover_letter_path": str(letter_path),
    }

# ─── Validação de completude ──────────────────────────────────────────────────


def validate_resume_completeness(md_text: str, profile: dict) -> list[str]:
    """
    Verifica se o currículo gerado contém os campos essenciais do perfil.

    Como o padrão ATS agora filtra skills/certificações/projetos menos
    relevantes para a vaga (intencionalmente), esta validação NÃO exige que
    todo o conteúdo do perfil apareça — apenas a presença das seções
    estruturais obrigatórias e das skills mais relevantes.

    Returns:
        Lista de warnings (vazia se completo o suficiente).
    """
    warnings_list = []

    # Seções estruturais obrigatórias
    if "## Resumo Profissional" not in md_text:
        warnings_list.append("Seção 'Resumo Profissional' ausente")
    if "## Experiência Profissional" not in md_text:
        warnings_list.append("Seção 'Experiência Profissional' ausente")
    if "## Habilidades" not in md_text and "## Habilidades Técnicas" not in md_text:
        warnings_list.append("Seção 'Habilidades' ausente")

    # Pelo menos uma skill relevante deve aparecer
    skills = profile.get("skills", [])
    if skills and not any(s in md_text for s in skills[:MAX_SKILLS]):
        warnings_list.append("Nenhuma skill do perfil encontrada no currículo")

    # Idiomas (se informados) — aceita string (split por vírgula) ou lista
    langs = profile.get("languages", "")
    if langs:
        lang_items = langs.split(",") if isinstance(langs, str) else list(langs)
        if not any(str(lang).strip() in md_text for lang in lang_items):
            warnings_list.append("Idiomas não encontrados no currículo")

    # Experience_raw vs md
    exp_raw = profile.get("experience_raw", "")
    if exp_raw and len(exp_raw) > 50:
        # Verifica se pelo menos 3 palavras significativas do raw aparecem
        words = set(re.findall(r'\b[A-Z][a-z]{3,}\b', exp_raw[:500]))
        matched = sum(1 for w in words if w.lower() in md_text.lower())
        if matched < max(3, len(words) // 3):
            warnings_list.append(f"Conteúdo de experiência parece incompleto ({matched}/{len(words)} termos)")

    return warnings_list
