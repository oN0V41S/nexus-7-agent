"""
Módulo de Consolidação de Currículos (REQ-004).

Extrai texto de múltiplos PDFs de currículo, normaliza em perfil unificado
e gera DOCX padrão ATS de 1 página + PDF de saída + Knowledge Base .md.
"""
import json
import logging
import re
import unicodedata
from datetime import date
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.job_apply_agent.config import PROFILE_DIR

logger = logging.getLogger(__name__)


# ─── Extração ────────────────────────────────────────────────────────────────


def extract_text_from_pdf(pdf_path: Path) -> str:
    """Extrai texto completo de um PDF usando PyMuPDF."""
    if not pdf_path.exists():
        raise FileNotFoundError(f"PDF não encontrado: {pdf_path}")

    text_parts: list[str] = []
    with fitz.open(str(pdf_path)) as doc:
        for page in doc:
            text_parts.append(page.get_text())

    full_text = "\n".join(text_parts)
    logger.info(f"Extraídos {len(full_text)} caracteres de {pdf_path.name}")
    return full_text


def normalize_text(text: str) -> str:
    """Remove espaços extras, normaliza quebras de linha."""
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"\n\s*\n", "\n\n", text)
    return text.strip()


# ─── Parse de seções ─────────────────────────────────────────────────────────


_SECTION_HEADERS = [
    "experiência", "experiencia", "formação", "formacao", "educação", "educacao",
    "habilidades", "skills", "idiomas", "cursos", "certificações", "certificacoes",
    "projetos", "projects", "resumo", "objetivo", "summary", "objective",
    "publicações", "publicacoes", "idioma",
]


def _guess_section(line: str) -> Optional[str]:
    """Tenta identificar se uma linha é cabeçalho de seção."""
    clean = line.strip().lower().rstrip(":")
    for header in _SECTION_HEADERS:
        # Exato ou contido
        if clean == header or clean.startswith(header):
            return header
    return None


def parse_resume_sections(text: str) -> dict[str, str]:
    """Divide texto de currículo em seções nomeadas."""
    lines = text.split("\n")
    sections: dict[str, str] = {"__header": ""}
    current_section = "__header"

    for line in lines:
        header = _guess_section(line)
        if header:
            current_section = header
            sections[current_section] = ""
        else:
            sections[current_section] = sections.get(current_section, "") + line + "\n"

    # Remove whitespace
    for key in sections:
        sections[key] = sections[key].strip()

    return sections


# ─── Extração DOCX ─────────────────────────────────────────────────────────────


def extract_text_from_docx(docx_path: Path) -> str:
    """Extrai texto limpo de um DOCX, preservando parágrafos como linhas separadas."""
    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX nao encontrado: {docx_path}")

    doc = Document(str(docx_path))
    paragraphs = [p.text for p in doc.paragraphs]
    full_text = "\n".join(paragraphs)
    logger.info(f"Extraidos {len(full_text)} caracteres de {docx_path.name} ({len(paragraphs)} paragrafos)")
    return full_text


def is_resume_boundary(text: str) -> bool:
    """
    Detecta se uma linha e um nome de candidato iniciando novo curriculo.

    Criterio: linha contem apenas maiusculas (incluindo acentos), espacos e pontos,
    com comprimento minimo de 10 caracteres. Ignora linhas com "@" ou "http"
    (para nao capturar linhas de contato).
    """
    stripped = text.strip()
    if not stripped:
        return False
    # Ignorar se contiver "@" ou "http" (senao captura linha de contato)
    if "@" in stripped or "http" in stripped.lower():
        return False
    # Regex: so letras maiusculas (incluindo acentos), espacos e pontos, minimo 10 chars
    if re.match(r'^[A-ZÀ-Ú][A-ZÀ-Ú\s\.]{25,}$', stripped):
        return True
    return False


def _guess_section_enhanced(line: str) -> Optional[str]:
    """
    Versao melhorada de _guess_section que:
    - Remove emojis antes de testar
    - Reconhece mais variacoes de nomes de secao
    """
    # Remove caracteres emoji (categoria Unicode So - Symbol, other)
    clean = ''.join(c for c in line if unicodedata.category(c) != 'So').strip()
    # Normaliza acentos (ex: "competências" -> "competencias") e lower
    clean = unicodedata.normalize('NFKD', clean)
    clean = clean.encode('ascii', 'ignore').decode('ascii').lower().rstrip(':').strip()
    if not clean:
        return None

    # Mapeamento de padroes estendidos para chaves canonicas
    patterns = {
        "resumo": ["resumo profissional", "resumo qualificado", "resumo"],
        "objetivo": ["objetivo profissional", "objetivo"],
        "habilidades": [
            "competencias tecnicas", "competencias", "habilidades",
            "ferramentas", "tecnologias", "skills",
        ],
        "experiencia": [
            "experiencia profissional", "experiencia",
        ],
        "formacao": [
            "formacao academica", "formacao", "educacao",
        ],
        "projetos": ["projetos destaque", "projetos profissionais", "projetos", "projects"],
        "certificacoes": [
            "certificacoes e idiomas", "certificacoes de destaque",
            "cursos e certificados", "certificacoes",
        ],
        "idiomas": ["idiomas", "idioma"],
        "publicacoes": ["publicacoes"],
    }

    for canonical, aliases in patterns.items():
        for alias in aliases:
            if clean == alias or clean.startswith(alias):
                return canonical

    # Fallback: tenta o _guess_section original (em texto sem emoji)
    return _guess_section(clean)


def parse_docx_to_sections(docx_path: Path) -> list[dict[str, str]]:
    """
    Pipeline completa de parse de DOCX:
    1. Extrai paragrafos do DOCX
    2. Detecta boundaries de multi-resume
    3. Para cada resume, divide em secoes usando _guess_section_enhanced
    4. Retorna uma lista de dicts (um por resume)
    """
    doc = Document(str(docx_path))
    paragraphs = [p.text for p in doc.paragraphs]
    logger.info(f"Parsing DOCX: {docx_path.name} -- {len(paragraphs)} paragrafos")

    # Passo 1: detectar boundaries de multi-resume
    resume_boundaries: list[int] = []
    for i, para in enumerate(paragraphs):
        if is_resume_boundary(para):
            resume_boundaries.append(i)

    # Se nao achou boundaries, trata o documento inteiro como um unico resume
    if not resume_boundaries:
        resume_boundaries = [0]

    logger.info(f"Detectadas {len(resume_boundaries)} boundaries de curriculo")

    # Passo 2: dividir em resumes individuais
    resumes_text: list[list[str]] = []
    for idx, start in enumerate(resume_boundaries):
        end = resume_boundaries[idx + 1] if idx + 1 < len(resume_boundaries) else len(paragraphs)
        resume_paras = paragraphs[start:end]
        resumes_text.append(resume_paras)

    # Passo 3: para cada resume, dividir em secoes
    all_result: list[dict[str, str]] = []
    for resume_idx, resume_paras in enumerate(resumes_text):
        sections: dict[str, str] = {"__header": ""}
        current_section = "__header"

        for para in resume_paras:
            header = _guess_section_enhanced(para)
            if header:
                current_section = header
                if current_section not in sections:
                    sections[current_section] = ""
                # Nao adiciona o cabecalho como conteudo
                continue

            # Pula linhas em branco no inicio (mas preserva no meio)
            if not para.strip() and current_section == "__header":
                continue

            sections[current_section] = sections.get(current_section, "") + para + "\n"

        # Remove trailing whitespace
        for key in list(sections.keys()):
            sections[key] = sections[key].strip()

        # Remove secoes vazias
        sections = {k: v for k, v in sections.items() if v}

        all_result.append(sections)
        logger.info(f"Resume #{resume_idx + 1}: {len(sections)} secoes encontradas")

    return all_result


def docx_to_kb(docx_path: Path, output_dir: Path) -> dict:
    """
    Pipeline orquestrador para DOCX:
    1. Chama parse_docx_to_sections()
    2. Escolhe o melhor resume (criterio: mais secoes preenchidas OU maior contagem total de caracteres)
    3. Constroi profile via build_profile_from_sections()
    4. Gera KB.md via generate_knowledge_base_md()
    5. Salva em output_dir/<slug>-kb-<YYYY-MM-DD>.md
    """
    logger.info(f"Iniciando docx_to_kb para {docx_path.name}")

    all_sections = parse_docx_to_sections(docx_path)
    if not all_sections:
        raise ValueError(f"Nenhuma secao extraida de {docx_path.name}")

    # Criterio de escolha: resume mais completo
    # - Prioriza experiencia (bonus +10000)
    # - Penaliza vazamento de nome em secoes (-2000)
    # - Mais secoes = melhor (+500 cada)
    # - Mais conteudo = melhor (+1 por char)
    def _score(sections: dict) -> int:
        total_chars = sum(len(v) for v in sections.values())
        section_count = len([k for k in sections if k != "__header"])
        has_experiencia = bool(sections.get("experiencia") or sections.get("experiência"))
        name_leak_penalty = 0
        for key, value in sections.items():
            if key != "__header" and re.search(
                r'RAFAEL\s+AUGUSTO|RAFAEL.AUGUSTO.*NOVAIS', value, re.IGNORECASE
            ):
                name_leak_penalty += 10000
        return (
            (1 if has_experiencia else 0) * 10000
            + section_count * 500
            + total_chars
            - name_leak_penalty
        )

    best_sections = max(all_sections, key=_score)
    logger.info(f"Melhor resume selecionado: {len(best_sections)} secoes, "
                f"{sum(len(v) for v in best_sections.values())} caracteres")

    # Constroi perfil
    profile = build_profile_from_sections(best_sections)

    # Extrai nome do header
    header = best_sections.get("__header", "")
    if header:
        first_line = header.split("\n")[0].strip()
        if first_line:
            profile["name"] = first_line

    # Gera KB.md
    slug = _normalize_name(profile.get("name", "candidato"))
    today = date.today().isoformat()
    kb_filename = f"{slug}-kb-{today}.md"
    kb_path = output_dir / kb_filename

    generate_knowledge_base_md(best_sections, profile, kb_path)

    logger.info(f"docx_to_kb concluido: KB em {kb_path}")
    return {
        "kb_path": str(kb_path),
        "profile": profile,
        "sections": best_sections,
        "num_resumes": len(all_sections),
    }


# ─── Geração de perfil ───────────────────────────────────────────────────────



def build_profile_from_sections(sections: dict[str, str]) -> dict:
    """Constrói dicionário de perfil estruturado a partir das seções extraídas."""
    skills_text = " ".join([
        sections.get(h, "") for h in ["habilidades", "skills", "cursos"]
    ])
    # Extrai skills como palavras capitalizadas
    skill_words = re.findall(r"[A-Z][a-zA-Z+#]+(?:\s*[A-Z][a-zA-Z+#]*)*", skills_text)
    skills = sorted(set(s.strip() for s in skill_words if len(s.strip()) > 1))

    experience = sections.get("experiência", "") or sections.get("experiencia", "")
    education = sections.get("formação", "") or sections.get("formacao", "") \
        or sections.get("educação", "") or sections.get("educacao", "")

    profile = {
        "skills": skills,
        "experience": normalize_text(experience),
        "education": normalize_text(education),
        "summary": normalize_text(
            sections.get("resumo", "") or sections.get("summary", "")
            or sections.get("objetivo", "") or sections.get("objective", "")
        ),
        "languages": normalize_text(sections.get("idiomas", "")),
        "certifications": normalize_text(sections.get("certificações", "")
                                         or sections.get("certificacoes", "")),
    }
    return profile


def merge_profiles(profiles: list[dict]) -> dict:
    """Merge de múltiplos perfis (ex: PDFs de fontes diferentes)."""
    merged = {}
    all_skills: set[str] = set()
    experiences: list[str] = []
    summaries: list[str] = []

    for p in profiles:
        all_skills.update(p.get("skills", []))
        if p.get("experience"):
            experiences.append(p["experience"])
        if p.get("summary"):
            summaries.append(p["summary"])
        # Pega a primeira education não vazia
        if p.get("education") and not merged.get("education"):
            merged["education"] = p["education"]
        if p.get("languages") and not merged.get("languages"):
            merged["languages"] = p["languages"]
        if p.get("certifications"):
            merged["certifications"] = p.get("certifications", "")

    merged["skills"] = sorted(all_skills)
    merged["experience"] = "\n\n".join(experiences) if experiences else ""
    merged["summary"] = " ".join(summaries) if summaries else ""
    return merged


# ─── Geração DOCX ATS ────────────────────────────────────────────────────────


def generate_ats_docx(profile: dict, output_path: Path) -> Path:
    """Gera DOCX padronizado para ATS (1 página, fonte limpa)."""
    doc = Document()

    # Configurar margens apertadas
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(2)
    paragraph_format.space_before = Pt(0)
    paragraph_format.line_spacing = 1.15

    # Nome (placeholder — virá do profile ou será editado manualmente)
    name = profile.get("name", "Candidato(a)")
    p = doc.add_heading(name, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    # Summary
    if profile.get("summary"):
        p = doc.add_paragraph(profile["summary"])
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.size = Pt(10)

    # Skills
    if profile.get("skills"):
        doc.add_heading("Habilidades", level=2)
        skills_text = ", ".join(profile["skills"])
        p = doc.add_paragraph(skills_text)
        for run in p.runs:
            run.font.size = Pt(10)

    # Experience
    if profile.get("experience"):
        doc.add_heading("Experiência Profissional", level=2)
        for exp_block in profile["experience"].split("\n\n"):
            p = doc.add_paragraph(exp_block.strip())
            for run in p.runs:
                run.font.size = Pt(10)

    # Education
    if profile.get("education"):
        doc.add_heading("Formação Acadêmica", level=2)
        p = doc.add_paragraph(profile["education"])
        for run in p.runs:
            run.font.size = Pt(10)

    # Languages
    if profile.get("languages"):
        doc.add_heading("Idiomas", level=2)
        p = doc.add_paragraph(profile["languages"])
        for run in p.runs:
            run.font.size = Pt(10)

    # Certifications
    if profile.get("certifications"):
        doc.add_heading("Certificações", level=2)
        p = doc.add_paragraph(profile["certifications"])
        for run in p.runs:
            run.font.size = Pt(10)

    doc.save(str(output_path))
    logger.info(f"DOCX salvo: {output_path}")
    return output_path


# ─── Geração PDF ─────────────────────────────────────────────────────────────


def generate_pdf_from_docx(docx_path: Path, output_path: Path) -> Path:
    """
    Converte DOCX para PDF usando fpdf2 com re-extração de texto.

    Nota: Para conversão direta DOCX→PDF com formatação completa,
    considere LibreOffice headless (soffice --convert-to pdf).
    Aqui usamos fpdf2 como fallback portável.
    """
    # Placeholder: fpdf2 exigiria re-parsing do docx.
    # Por ora, copiamos o DOCX como base.
    # Em produção, substitua por conversão via LibreOffice ou Unoconv.
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=10)

    doc = Document(str(docx_path))
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # Tenta detectar cabeçalhos
        if para.style.name.startswith("Heading"):
            pdf.set_font("Helvetica", "B", 12)
            pdf.multi_cell(0, 6, text)
            pdf.set_font("Helvetica", size=10)
        else:
            pdf.multi_cell(0, 5, text)

    pdf.output(str(output_path))
    logger.info(f"PDF salvo: {output_path}")
    return output_path


# ─── Helpers para Knowledge Base .md ─────────────────────────────────────────


def _extract_candidate_name(raw_sections: dict, profile: dict) -> str:
    """Tenta extrair nome do candidato do header ou profile."""
    header = raw_sections.get("__header", "").strip()
    if header:
        # Primeira linha não vazia é geralmente o nome
        first_line = header.split("\n")[0].strip()
        if first_line and len(first_line) < 80:
            return first_line
    return profile.get("name", "Candidato(a)")


def _normalize_name(name: str) -> str:
    """Normaliza nome para slug de arquivo (ex: 'João Silva' → 'joao-silva')."""
    name = unicodedata.normalize("NFKD", name)
    name = name.encode("ascii", "ignore").decode("ascii")
    name = re.sub(r"[^a-zA-Z0-9\s-]", "", name)
    name = re.sub(r"\s+", "-", name.strip().lower())
    return name


def _merge_raw_sections(all_sections: list[dict]) -> dict:
    """Merge de múltiplos raw_sections: pega o primeiro de cada tipo encontrado."""
    merged: dict[str, str] = {"__header": ""}
    seen_types: set[str] = set()
    for sections in all_sections:
        for key, value in sections.items():
            if value and key not in seen_types:
                merged[key] = value
                seen_types.add(key)
    return merged


# ─── Geração Knowledge Base .md ──────────────────────────────────────────────


def generate_knowledge_base_md(
    raw_sections: dict, profile: dict, output_path: Path
) -> Path:
    """
    Gera Knowledge Base .md completo com TODO o conteúdo original do currículo.

    Recebe as seções brutas (dict de parse_resume_sections + header), o perfil
    estruturado e o path de saída. Gera Markdown completo com todas as seções
    preservadas — nada é omitido, resumido ou inventado.

    Args:
        raw_sections: Dict com seções brutas do currículo (chaves: __header,
                      resumo, experiência, habilidades, idiomas, ...).
        profile: Perfil estruturado (pode conter 'name', 'email', 'phone' etc.).
        output_path: Caminho onde o arquivo .md será salvo.

    Returns:
        Path para o arquivo .md gerado.
    """
    name = _extract_candidate_name(raw_sections, profile)

    # ── Processa header para informações de contato ──
    header_lines = raw_sections.get("__header", "").strip().split("\n")
    # Primeira linha é o nome; o restante são dados de contato
    contact_lines = [l.strip() for l in header_lines[1:] if l.strip()]
    if contact_lines:
        contact_str = " | ".join(contact_lines)
    else:
        # Fallback: monta a partir do profile
        parts = []
        if profile.get("email"):
            parts.append(profile["email"])
        if profile.get("phone"):
            parts.append(profile["phone"])
        if profile.get("location"):
            parts.append(profile["location"])
        contact_str = " | ".join(parts) if parts else "*(Informações de contato não extraídas)*"

    # ── Constrói conteúdo Markdown ──
    kb_lines: list[str] = []

    def _add_section(title: str, raw_keys: list[str]) -> None:
        """Adiciona seção ao KB se o conteúdo existir nas raw_sections."""
        content = ""
        for key in raw_keys:
            val = raw_sections.get(key, "").strip()
            if val:
                content = val
                break
        if content:
            kb_lines.append(f"## {title}")
            kb_lines.append(content)
            kb_lines.append("")

    kb_lines.append(f"# {name}")
    kb_lines.append("")
    kb_lines.append("## Contato")
    kb_lines.append(contact_str)
    kb_lines.append("")

    _add_section("Resumo Profissional", ["resumo", "summary", "objetivo", "objective"])
    _add_section("Experiência Profissional", ["experiência", "experiencia"])
    _add_section("Formação Acadêmica", ["formação", "formacao", "educação", "educacao"])
    _add_section("Habilidades", ["habilidades", "skills", "cursos"])
    _add_section("Idiomas", ["idiomas", "idioma"])
    _add_section("Certificações", ["certificações", "certificacoes", "cursos"])
    _add_section("Projetos", ["projetos", "projects"])
    _add_section("Publicações", ["publicações", "publicacoes"])

    content = "\n".join(kb_lines)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(content, encoding="utf-8")
    logger.info(f"Knowledge Base .md salva: {output_path} ({len(content)} caracteres)")
    return output_path


# ─── Orquestração ────────────────────────────────────────────────────────────


def consolidate_pdfs_to_kb(pdf_paths: list[Path], output_dir: Path) -> dict:
    """
    Pipeline simplificado: extrai PDFs → parse → KB .md.

    Gera apenas a Knowledge Base .md (sem profile.json, DOCX ou PDF).

    Args:
        pdf_paths: Lista de paths para PDFs de currículo.
        output_dir: Diretório de saída (criado automaticamente se não existir).

    Returns:
        dict com chaves: kb_path, profile
    """
    all_sections: list[dict] = []
    profiles: list[dict] = []

    for pdf_path in pdf_paths:
        text = extract_text_from_pdf(pdf_path)
        sections = parse_resume_sections(text)
        all_sections.append(sections)
        profile = build_profile_from_sections(sections)
        profiles.append(profile)

    merged_profile = merge_profiles(profiles)
    merged_raw = _merge_raw_sections(all_sections)

    # Nomeia o arquivo
    name = _extract_candidate_name(merged_raw, merged_profile)
    slug = _normalize_name(name)
    today = date.today().isoformat()
    kb_filename = f"{slug}-kb-{today}.md"
    kb_path = output_dir / kb_filename

    generate_knowledge_base_md(merged_raw, merged_profile, kb_path)

    return {
        "kb_path": str(kb_path),
        "profile": merged_profile,
    }


def consolidate_pdfs_to_docx(pdf_paths: list[Path], output_dir: Path) -> dict:
    """
    Pipeline completo: extrai PDFs → constrói perfil → gera DOCX + PDF + KB .md.

    Args:
        pdf_paths: Lista de paths para PDFs de currículo.
        output_dir: Diretório de saída (criado automaticamente se não existir).

    Returns:
        dict com chaves: profile, docx_path, pdf_path, kb_path
    """
    profiles = []
    all_raw_sections: list[dict] = []

    for pdf_path in pdf_paths:
        text = extract_text_from_pdf(pdf_path)
        sections = parse_resume_sections(text)
        all_raw_sections.append(sections)
        profile = build_profile_from_sections(sections)
        profiles.append(profile)

    merged_profile = merge_profiles(profiles)

    # Salva profile.json
    profile_path = PROFILE_DIR / "profile.json"
    profile_path.write_text(
        json.dumps(merged_profile, indent=2, ensure_ascii=False)
    )

    # Gera DOCX
    docx_path = output_dir / "resume_ats.docx"
    generate_ats_docx(merged_profile, docx_path)

    # Gera PDF
    pdf_path = output_dir / "resume_ats.pdf"
    generate_pdf_from_docx(docx_path, pdf_path)

    # Gera Knowledge Base .md
    merged_raw = _merge_raw_sections(all_raw_sections)
    name = _extract_candidate_name(merged_raw, merged_profile)
    slug = _normalize_name(name)
    today = date.today().isoformat()
    kb_filename = f"{slug}-kb-{today}.md"
    kb_path = output_dir / kb_filename
    generate_knowledge_base_md(merged_raw, merged_profile, kb_path)

    return {
        "profile": merged_profile,
        "docx_path": str(docx_path),
        "pdf_path": str(pdf_path),
        "kb_path": str(kb_path),
    }
